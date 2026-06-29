"""
تصدير صكّ الحكم إلى ملفّ Word (DOCX) منسّقٍ بالعربية (RTL) — منتَجٌ للقاضي/المحامي.
يبني: ترويسةً وبيانات الدعوى، ثم الوقائع والأسباب والمنطوق، ثم سلسلة الاستدلال
والإسنادات، وتذييلَ التنبيه. مدخله حمولةُ الحكم القادمة من الواجهة.
"""
from __future__ import annotations

from io import BytesIO

from . import settings


def _rtl_paragraph(p) -> None:
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))


def _rtl_run(r) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    rPr = r._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


def build_ruling_docx(payload: dict) -> bytes:
    """يبني صكّ الحكم DOCX (RTL) من حمولة الواجهة ويُرجعه bytes."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)

    def para(text, size=12, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        _rtl_paragraph(p)
        for line in (text or "—").split("\n"):
            r = p.add_run(line)
            r.font.size = Pt(size)
            r.bold = bold
            if color:
                r.font.color.rgb = RGBColor(*color)
            _rtl_run(r)
            r.add_break()
        return p

    def heading(text):
        para(text, size=14, bold=True, color=(0x26, 0x21, 0x5C), space_after=3)

    meta = payload.get("meta") or {}
    GOLD = (0x99, 0x3C, 0x1D)

    # الترويسة
    para("المملكة العربية السعودية — محاكاة المحكمة التجارية", size=11, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(f"صكّ {payload.get('kind', 'حكم')} — منصّة «بيّن»", size=16, bold=True,
         color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para("⚠️ محاكاةٌ تدريبية فقط — ليست حكماً قضائياً حقيقياً ولا استشارةً قانونية.",
         size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x35, 0x56), space_after=10)

    # بيانات الدعوى
    heading("بيانات الدعوى")
    info = [("رقم الدعوى", meta.get("case_id", "—")), ("النوع", meta.get("case_type", "تجاري")),
            ("قيمة المطالبة", f"{meta.get('value', '—')} ر.س"), ("تاريخ القيد", meta.get("filing", "—"))]
    if payload.get("confidence"):
        info.append(("ثقة المحرّك", payload["confidence"]))
    if payload.get("route"):
        appl = "قابل للاستئناف" if payload.get("appealable") else "نهائي"
        info.append(("قابلية الاعتراض", f"{appl} — {payload['route']}"))
    para("   |   ".join(f"{k}: {v}" for k, v in info), size=11, space_after=10)

    if payload.get("blocked"):
        para("⛔ حُجِب النطق: حكمٌ مبنيٌّ على إسنادٍ مختلق أو مخالفةٍ جوهرية — لا يُعتدّ به.",
             size=11, bold=True, color=(0xB2, 0x3A, 0x5B), space_after=8)

    heading("أولاً: الوقائع")
    para(payload.get("facts"))
    heading("ثانياً: الأسباب (التسبيب)")
    para(payload.get("reasons"))
    heading("ثالثاً: المنطوق")
    para(payload.get("operative"), bold=True)

    # سلسلة الاستدلال
    chain = payload.get("chain") or []
    if chain:
        heading("سلسلة الاستدلال القضائي")
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        t.alignment = 2  # RTL table
        _cell(t.rows[0].cells[0], "العملية", bold=True)
        _cell(t.rows[0].cells[1], "الحالة", bold=True)
        for c in chain:
            row = t.add_row().cells
            _cell(row[0], c.get("label", ""))
            issues = "؛ ".join(c.get("issues", []) or [])
            _cell(row[1], ("✓ سليمة" if c.get("ok") else "⚠ " + issues) if not issues or c.get("ok") else "⚠ " + issues)
        doc.add_paragraph()

    # الإسنادات
    cites = payload.get("citations") or []
    if cites:
        heading("الإسنادات النظامية")
        t = doc.add_table(rows=1, cols=3)
        t.style = "Table Grid"
        t.alignment = 2
        for i, h in enumerate(("الادعاء", "المصدر", "المرجع")):
            _cell(t.rows[0].cells[i], h, bold=True)
        for c in cites:
            row = t.add_row().cells
            _cell(row[0], c.get("claim", ""))
            _cell(row[1], c.get("tool", ""))
            _cell(row[2], c.get("ref", ""))
        doc.add_paragraph()

    flags = [f for f in (payload.get("flags") or []) if f]
    if flags:
        heading("ملاحظات التأصيل")
        for f in flags:
            para("• " + f, size=10, color=(0x99, 0x3C, 0x1D), space_after=2)

    para("\n" + settings.DISCLAIMER, size=9, align=WD_ALIGN_PARAGRAPH.CENTER,
         color=(0x88, 0x87, 0x80), space_after=0)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _cell(cell, text, bold=False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _rtl_paragraph(p)
    r = p.add_run(text or "—")
    r.font.size = Pt(10.5)
    r.bold = bold
    _rtl_run(r)
